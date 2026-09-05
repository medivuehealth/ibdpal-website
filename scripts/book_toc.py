"""Chapter table of contents for the book compiler."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

STYLE_FRONT = "Book Front Matter"
STYLE_SECTION = "Book Section"
STYLE_KINDLE_TOC_1 = "Kindle TOC 1"
STYLE_KINDLE_TOC_2 = "Kindle TOC 2"

# Print TOC uses Word auto-TOC levels 1–2 only (parts, chapters, major back matter).
PRINT_TOC_FIELD = r'TOC \o "1-2" \h \z \u'

FORBIDDEN_TOC_SUBSTRINGS = (
    "macronutrients",
    "micronutrients",
    "adult female",
    "white rice",
    "during a flare",
    "during remission",
)

PRINT_BACK_MATTER = (
    "Conclusion: Eating With IBD Over Time",
    "About the Author",
    "Emergency and Red-Flag Symptoms",
    "Glossary",
    "Abbreviations",
    "Food Index",
    "Index",
    "Illustration Credits",
)


def ensure_book_styles(doc: Document) -> None:
    """Non-TOC paragraph styles for front matter and in-chapter sections."""
    styles = doc.styles

    def _make(name: str, base: str, size: int, bold: bool, space_before: int = 12) -> None:
        try:
            st = styles[name]
        except KeyError:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = styles[base]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(6)

    _make(STYLE_FRONT, "Normal", 14, True, space_before=18)
    _make(STYLE_SECTION, "Normal", 12, True, space_before=10)
    _ensure_kindle_toc_styles(doc)


def _ensure_kindle_toc_styles(doc: Document) -> None:
    for name, size, indent in ((STYLE_KINDLE_TOC_1, 11, 0), (STYLE_KINDLE_TOC_2, 10.5, 18)):
        try:
            st = doc.styles[name]
        except KeyError:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = doc.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        pf = st.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        if indent:
            pf.left_indent = Pt(indent)


def _ensure_word_builtin_toc_styles(doc: Document) -> None:
    """Tab leaders on built-in TOC styles so titles and page numbers do not concatenate."""
    for name in ("toc 1", "toc 2"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = "Calibri"
        st.font.size = Pt(11 if name == "toc 1" else 10.5)
        pf = st.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.tab_stops.clear_all()
        pf.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def _insert_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {PRINT_TOC_FIELD} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_section_heading(doc: Document, text: str, *, level: int = 2, kindle: bool = False) -> None:
    """In-chapter subsection: Book Section style (excluded from print and Kindle NCX)."""
    del level, kindle
    ensure_book_styles(doc)
    doc.add_paragraph(text, style=STYLE_SECTION)


def add_front_matter_heading(doc: Document, text: str, *, kindle: bool = False) -> None:
    """Front matter titles excluded from print auto-TOC and Kindle NCX."""
    del kindle
    ensure_book_styles(doc)
    doc.add_paragraph(text, style=STYLE_FRONT)


class KindleNavRegistry:
    """Bookmarks for chapter-level Kindle navigation (parts, chapters, back matter)."""

    def __init__(self) -> None:
        self._next_id = 1

    def attach_bookmark(self, paragraph, name: str) -> None:
        bookmark_id = self._next_id
        self._next_id += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    @staticmethod
    def part_name(num: int) -> str:
        return f"bm_part_{num}"

    @staticmethod
    def chapter_name(num: int) -> str:
        return f"bm_chapter_{num}"

    @staticmethod
    def back_matter_name(title: str) -> str:
        slug = re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_")
        return f"bm_back_{slug}"


def add_paragraph_hyperlink(paragraph, text: str, bookmark: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_back_matter_heading(
    doc: Document,
    text: str,
    *,
    kindle: bool = False,
    nav: KindleNavRegistry | None = None,
) -> None:
    """Major back matter: Heading 1 in print (included in auto-TOC)."""
    del kindle
    paragraph = doc.add_heading(text, level=1)
    if nav is not None:
        nav.attach_bookmark(paragraph, KindleNavRegistry.back_matter_name(text))


def add_book_heading(
    doc: Document,
    text: str,
    *,
    part: bool = False,
    kindle: bool = False,
    nav: KindleNavRegistry | None = None,
    chapter_num: int | None = None,
    part_num: int | None = None,
) -> None:
    """Parts and chapters use built-in Heading 1/2 so Word auto-TOC resolves page numbers."""
    del kindle
    paragraph = doc.add_heading(text, level=1 if part else 2)
    if nav is None:
        return
    if part and part_num is not None:
        nav.attach_bookmark(paragraph, KindleNavRegistry.part_name(part_num))
    elif not part and chapter_num is not None:
        nav.attach_bookmark(paragraph, KindleNavRegistry.chapter_name(chapter_num))


def build_kindle_toc_entries(
    chapters: list[dict],
    parts: list[dict],
) -> list[tuple[str, int, str]]:
    entries: list[tuple[str, int, str]] = []
    for part in parts:
        pnum = part["num"]
        title = f"Part {pnum}: {part['title']}"
        entries.append((title, 1, KindleNavRegistry.part_name(pnum)))
        for ch in chapters:
            if ch["part"] == pnum:
                ctitle = f"Chapter {ch['num']}: {ch['title']}"
                entries.append((ctitle, 2, KindleNavRegistry.chapter_name(ch["num"])))
    for back_title in PRINT_BACK_MATTER:
        entries.append(
            (back_title, 1, KindleNavRegistry.back_matter_name(back_title)),
        )
    return entries


def add_chapter_table_of_contents(
    doc: Document,
    *,
    kindle: bool = False,
    chapters: list[dict] | None = None,
    parts: list[dict] | None = None,
    nav: KindleNavRegistry | None = None,
) -> None:
    """Print: Word auto-TOC (Heading 1–2). Kindle: hyperlinked chapter nav without page numbers."""
    if chapters is None or parts is None:
        raise ValueError("TOC requires chapters and parts manifest data")

    if kindle:
        ensure_book_styles(doc)
        for title, level, bookmark in build_kindle_toc_entries(chapters, parts):
            style = STYLE_KINDLE_TOC_1 if level == 1 else STYLE_KINDLE_TOC_2
            paragraph = doc.add_paragraph(style=style)
            add_paragraph_hyperlink(paragraph, title, bookmark)
        return

    _ensure_word_builtin_toc_styles(doc)
    _insert_toc_field(doc.add_paragraph())


def verify_print_toc(docx_path: str | Path, *, kindle: bool = False) -> list[str]:
    """Return errors if TOC has broken fields, PAGEREF residue, or subsection leaks."""
    errors: list[str] = []
    doc = Document(str(docx_path))
    bad_pageref = False
    for p in doc.paragraphs:
        for el in p._element.iter():
            if el.tag.endswith("instrText") and el.text and "PAGEREF" in el.text:
                if "PAGEREF bm_" in el.text:
                    bad_pageref = True
        text = p.text.strip()
        if "Error" in text and "Reference" in text:
            errors.append(f"Broken reference in document: {text[:80]}")

    if kindle:
        if bad_pageref:
            errors.append("Kindle edition must not contain legacy bm_ PAGEREF fields.")
        has_auto_toc = False
        in_toc = False
        toc_lines: list[str] = []
        heading3_plus = 0
        for p in doc.paragraphs:
            for el in p._element.iter():
                if el.tag.endswith("instrText") and el.text and "TOC \\" in el.text:
                    has_auto_toc = True
            style_name = (p.style.name if p.style else "") or ""
            if style_name in ("Heading 3", "Heading 4", "Heading 5"):
                heading3_plus += 1
            text = p.text.strip()
            if text == "Table of Contents":
                in_toc = True
                continue
            if in_toc and text == "Start Here":
                break
            if in_toc and text:
                toc_lines.append(text)
                if re.search(r"\t\d{1,4}$", text):
                    errors.append(f"Kindle TOC has print page number: {text[:80]}")
        if has_auto_toc:
            errors.append('Kindle edition must not contain auto-TOC field (TOC \\o).')
        if len(toc_lines) > 80:
            errors.append(f"Kindle TOC has {len(toc_lines)} lines; expected ~68 chapter-level entries.")
        if heading3_plus:
            errors.append(
                f"Kindle edition has {heading3_plus} Heading 3+ paragraphs; subsections belong in Book Section style.",
            )
        for line in toc_lines:
            lower = line.lower()
            if lower.startswith(("part ", "chapter ")):
                continue
            if any(sub in lower for sub in FORBIDDEN_TOC_SUBSTRINGS):
                errors.append(f"Kindle TOC subsection leaked: {line[:80]}")
        return errors

    in_toc = False
    toc_lines: list[str] = []
    has_auto_toc = False

    for p in doc.paragraphs:
        for el in p._element.iter():
            if el.tag.endswith("instrText") and el.text and "TOC \\" in el.text:
                has_auto_toc = True

        text = p.text.strip()
        if text == "Table of Contents":
            in_toc = True
            continue
        if in_toc and text == "Start Here":
            break
        if in_toc and text:
            line = text.split("\t")[0]
            toc_lines.append(line.lower())
            if re.search(r"\d+Chapter\s+\d+", line, re.I):
                errors.append(f"TOC concatenation: {line[:80]}")
            if re.search(r"[a-z]\d{1,3}$", line.replace(" ", ""), re.I) and "chapter" not in line.lower()[:8]:
                errors.append(f"TOC title/page merged: {line[:80]}")

    if not has_auto_toc:
        errors.append('Print TOC missing auto-TOC field (TOC \\o "1-2").')
    if bad_pageref:
        errors.append("Print TOC still contains legacy bm_ PAGEREF fields.")
    if len(toc_lines) > 80:
        errors.append(f"TOC has {len(toc_lines)} lines; expected ~68 chapter-level entries.")
    for line in toc_lines:
        if line.startswith(("part ", "chapter ")):
            continue
        if any(sub in line for sub in FORBIDDEN_TOC_SUBSTRINGS):
            errors.append(f"TOC subsection leaked: {line[:80]}")
    return errors


def update_docx_fields(docx_path: str) -> bool:
    """Format TOC tab leaders and refresh page numbers via Word (Windows)."""
    path = str(docx_path)
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False

    wdLeaderDots = 2
    wdAlignTabRight = 2

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path)

        right_tab = float(doc.PageSetup.PageWidth) - float(doc.PageSetup.RightMargin)
        for style_name in ("TOC 1", "TOC 2"):
            try:
                style = doc.Styles(style_name)
                style.ParagraphFormat.TabStops.ClearAll()
                style.ParagraphFormat.TabStops.Add(
                    Position=right_tab,
                    Alignment=wdAlignTabRight,
                    Leader=wdLeaderDots,
                )
            except Exception:
                pass

        doc.Fields.Update()
        doc.Repaginate()
        for toc in doc.TablesOfContents:
            toc.Update()
        doc.Fields.Update()
        doc.Save()
        doc.Close(SaveChanges=True)
        return True
    except Exception:
        return False
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            word.Quit()
