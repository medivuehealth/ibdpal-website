"""Back-matter appendices: glossary, abbreviations, index, clinical checklist."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from book_manifest_data import CHAPTERS
from book_text_cleanup import finalize_text

GLOSSARY: list[tuple[str, str]] = [
    ("Active disease / flare", "Period of increased intestinal inflammation and/or symptoms requiring medical attention."),
    ("Calprotectin", "Stool marker of intestinal inflammation; used with symptoms and scopes."),
    ("EEN", "Exclusive enteral nutrition, formula-only nutrition used as induction therapy in selected patients, especially pediatric Crohn's disease."),
    ("Enteral nutrition", "Nutrition delivered through a tube or oral formula."),
    ("FODMAP", "Fermentable oligosaccharides, disaccharides, monosaccharides, and polyols, short-chain carbohydrates that may worsen gas and bloating in some people."),
    ("Low residue", "Diet pattern reducing fiber and indigestible material; temporary for some during flares."),
    ("Malabsorption", "Incomplete absorption of nutrients in the small intestine."),
    ("Remission", "Period of reduced or absent significant inflammation; symptoms may still occur."),
    ("RDN", "Registered dietitian nutritionist, credentialed nutrition professional."),
    ("Stricture", "Narrowing of the intestine that may require texture modification."),
]

ABBREVIATIONS: list[tuple[str, str]] = [
    ("ACG", "American College of Gastroenterology"),
    ("AGA", "American Gastroenterological Association"),
    ("CD", "Crohn's disease"),
    ("CDED", "Crohn's Disease Exclusion Diet"),
    ("CRP", "C-reactive protein"),
    ("DRI", "Dietary Reference Intakes (NIH)"),
    ("ECCO", "European Crohn's and Colitis Organisation"),
    ("EEN", "Exclusive enteral nutrition"),
    ("ESPGHAN", "European Society for Paediatric Gastroenterology, Hepatology and Nutrition"),
    ("FODMAP", "Fermentable oligosaccharides, disaccharides, monosaccharides, and polyols"),
    ("GI", "Gastroenterologist or gastrointestinal"),
    ("IBD", "Inflammatory bowel disease"),
    ("NASPGHAN", "North American Society for Pediatric Gastroenterology, Hepatology and Nutrition"),
    ("RDN", "Registered dietitian nutritionist"),
    ("SCD", "Specific Carbohydrate Diet"),
    ("UC", "Ulcerative colitis"),
]

RED_FLAGS = [
    "Severe abdominal pain that does not improve or feels different from your usual flare",
    "Repeated vomiting or inability to keep fluids down",
    "Heavy rectal bleeding or dizziness with bleeding",
    "High fever with worsening abdominal symptoms",
    "Signs of dehydration: dark urine, dizziness, rapid heartbeat, confusion",
    "Rapid unintentional weight loss",
    "Suspected bowel obstruction: bloating, no gas or stool, severe pain, vomiting",
]

CLINICAL_REVIEW_CHECKLIST = [
    "Exclusive enteral nutrition indications and wording",
    "Stricture and fiber guidance, no unsafe blanket recommendations",
    "Low-residue / low-fiber advice qualified by anatomy and disease state",
    "Iron, anemia, B12, vitamin D, calcium, zinc, magnesium, folate guidance",
    "Electrolyte and dehydration recommendations",
    "Supplement dose language, no universal prescribing",
    "Pregnancy and lactation nutrition",
    "Pediatric growth and EEN",
    "Ostomy and short bowel hydration/salt guidance",
    "Post-surgery reintroduction staging",
    "Medication, nutrition interactions (steroids, methotrexate, biologics)",
    "FODMAP, CDED, SCD, Mediterranean, evidence strength accurately labeled",
    "Symptom vs inflammation distinction preserved",
    "Emergency / red-flag symptoms appropriately emphasized",
]


def build_index_entries() -> list[tuple[str, str]]:
    entries: list[tuple[str, str]] = []
    for ch in CHAPTERS:
        entries.append((ch["title"], f"Chapter {ch['num']}"))
    nutrients = [
        "Iron deficiency", "Anemia", "Vitamin B12", "Vitamin D", "Calcium",
        "Folate", "Zinc", "Magnesium", "Potassium", "Protein", "Fiber", "Hydration",
        "Electrolytes", "Omega-3", "Probiotics", "Supplements",
    ]
    for n in nutrients:
        entries.append((n, "Part III; see Table of Contents"))
    foods = [
        "Banana", "Rice", "Dal", "Lentils", "Chicken", "Salmon", "Eggs", "Oatmeal",
        "Potato", "Yogurt", "Coffee", "Congee", "Paneer", "Tofu", "Plantain",
    ]
    for f in foods:
        entries.append((f, "Part VI, Food Reference Library"))
    topics = [
        ("FODMAP", "Chapters 23, 24"),
        ("Exclusive enteral nutrition", "Chapters 28, 31"),
        ("Ostomy nutrition", "Chapter 48"),
        ("Pregnancy", "Chapter 46"),
        ("Food tracking", "Chapter 49"),
        ("Personal nutrition plan", "Chapter 51"),
    ]
    entries.extend(topics)
    entries.sort(key=lambda x: x[0].lower())
    return entries


def add_red_flags_page(doc: Document, *, kindle: bool = False, nav=None) -> None:
    from book_toc import add_back_matter_heading

    doc.add_page_break()
    add_back_matter_heading(
        doc, "Emergency and Red-Flag Symptoms", kindle=kindle, nav=nav,
    )
    doc.add_paragraph(
        finalize_text(
            "Seek urgent medical care, do not rely on this book alone, if you experience:"
        )
    )
    for item in RED_FLAGS:
        doc.add_paragraph(item, style="List Bullet")


def add_glossary(doc: Document, *, kindle: bool = False, nav=None) -> None:
    from book_toc import add_back_matter_heading

    doc.add_page_break()
    add_back_matter_heading(doc, "Glossary", kindle=kindle, nav=nav)
    for term, definition in GLOSSARY:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(finalize_text(definition))


def add_abbreviations(doc: Document, *, kindle: bool = False, nav=None) -> None:
    from book_toc import add_back_matter_heading

    doc.add_page_break()
    add_back_matter_heading(doc, "Abbreviations", kindle=kindle, nav=nav)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Abbreviation"
    table.rows[0].cells[1].text = "Meaning"
    for abbr, meaning in ABBREVIATIONS:
        row = table.add_row().cells
        row[0].text = abbr
        row[1].text = finalize_text(meaning)


def add_index(doc: Document, *, kindle: bool = False, nav=None) -> None:
    from book_toc import add_back_matter_heading

    doc.add_page_break()
    add_back_matter_heading(doc, "Index", kindle=kindle, nav=nav)
    if kindle:
        doc.add_paragraph(
            "Use this index for quick lookup by chapter and part. Kindle navigation also "
            "lists parts and chapters in the device menu.",
            style="Intense Quote",
        )
    else:
        doc.add_paragraph(
            "Use this index for quick lookup. Page numbers refer to chapter locations; "
            "refresh final page numbers after layout.",
            style="Intense Quote",
        )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Topic"
    table.rows[0].cells[1].text = "See"
    for topic, loc in build_index_entries():
        row = table.add_row().cells
        row[0].text = topic
        row[1].text = loc


def add_clinical_review_checklist(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix: Clinical Review Checklist", level=1)
    doc.add_paragraph(
        "For gastroenterology and registered dietitian reviewers, mark any section needing "
        "correction, qualification, or removal before publication."
    )
    for item in CLINICAL_REVIEW_CHECKLIST:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")


def add_food_index(doc: Document, *, kindle: bool = False, nav=None) -> None:
    from book_toc import add_back_matter_heading

    doc.add_page_break()
    add_back_matter_heading(doc, "Food Index", kindle=kindle, nav=nav)
    food_chapters = [c for c in CHAPTERS if c["part"] == 6 and c["num"] <= 38]
    for ch in food_chapters:
        p = doc.add_paragraph()
        run = p.add_run(f"Chapter {ch['num']}: {ch['title']}")
        run.bold = True
        run.font.size = Pt(11)
        for src in ch.get("sources", []):
            slug = src.split("/")[-1].replace(".html", "").replace("-ibd", "").replace("-", " ")
            doc.add_paragraph(slug.title(), style="List Bullet 2")
