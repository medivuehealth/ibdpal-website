#!/usr/bin/env python3
"""Search a compiled DOCX for production placeholders before KDP upload."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = ROOT / "content" / "ibd-nutrition-book" / "Eating_With_IBD_Interior.docx"

REGRESSION_PHRASES = [
    "increasing calorie needs",
    "Surgery, resection, strictureplasty, ostomy, or pouch, can improve inflammation",
    "may affect absorption of iron, B12, and fat-soluble vitamins",
    "Request monitoring at routine visits",
    "Unintentional weight change is a clinical clue",
    "Track weight weekly at the same time of day",
    "Iron, B12, vitamin D, zinc, and calcium need periodic labs",
    "Absorption site for many nutrients in the upper GI tract",
]

SEARCH_TERMS = [
    "Error: Reference",
    "Reference source not found",
    "Error! Bookmark",
    "placeholder",
    "XXXXX",
    "TBD",
    "TODO",
    "insert diagram",
    "[Name]",
    "reviewed for medical",
    "Manuscript Compile Report",
    "compile report",
    "missing_sources",
    "estimated_pages",
    "BARCODE",
    "Unsplash License (free to use)",
    "click here",
    "read more",
    "see article",
    "related article",
    "complete nutrition article",
    "foundation basics",
    "editorial revision pipeline",
    "978-1-XXXXX",
    "pipeline",
    "manifest",
    "draft",
    "Figure 1.1",
    "Figure 1.2",
    "Figure 2.1",
    "Fig. 1.1",
    "Fig. 1.2",
    "Fig. 2.1",
    "as shown in Figure",
    "shown below",
    "shown above",
    "Foods during a Crohn",
    "Complete nutrition article",
    "What is Crohn",
    "Foundation basics",
    "learn more",
    "patient education materials",
]

FALSE_POSITIVE = [
    re.compile(r"extraintestinal manifest", re.I),
    re.compile(r"prompt clinical review", re.I),
    re.compile(r"not been formally clinically reviewed", re.I),
    re.compile(r"Illustration Credits", re.I),
    re.compile(r"unsplash\.com/license", re.I),
    re.compile(r"pexels\.com/license", re.I),
    re.compile(r"what is crohn", re.I),  # legitimate chapter prose
    re.compile(r"crohn'?s disease", re.I),
    re.compile(r"where can i read more", re.I),
    re.compile(r"helps you draft a personal plan", re.I),
    re.compile(r"^Figure 1\.1\. Digestive tract overview", re.I),
    re.compile(r"^Figure 1\.2\. Symptoms versus inflammation", re.I),
    re.compile(r"^Figure 2\.1\. Disease location and absorption", re.I),
]


MECHANICAL_DEFECTS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"References\s*\d+\.", re.I), "References glued to citation number"),
    (re.compile(r"Key Takeaways\s*[A-Za-z]"), "Key Takeaways glued to following text"),
    (
        re.compile(r"^\d+\.\s+.+\s\d+\.\s+(?:[A-Z\"'(]|National|NIH|USDA|Crohn|World|American)"),
        "Multiple citations in one paragraph",
    ),
    (
        re.compile(r"_______________\s+\S.{10,}_______________"),
        "Worksheet fields glued in one paragraph",
    ),
]


def scan_mechanical_defects(path: Path) -> list[tuple[int, str, str]]:
    doc = Document(str(path))
    hits: list[tuple[int, str, str]] = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, label in MECHANICAL_DEFECTS:
            if pattern.search(text):
                hits.append((i + 1, label, text[:160]))
                break
    return hits


def scan_docx(path: Path) -> list[tuple[int, str, str]]:
    doc = Document(str(path))
    hits: list[tuple[int, str, str]] = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        lower = text.lower()
        for term in REGRESSION_PHRASES + SEARCH_TERMS:
            if term.lower() in lower:
                if term not in REGRESSION_PHRASES and any(p.search(text) for p in FALSE_POSITIVE):
                    continue
                hits.append((i + 1, term, text[:160]))
                break
    return hits


def main() -> int:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX
    if not path.is_file():
        print(f"File not found: {path}")
        return 1
    hits = scan_docx(path)
    mechanical = scan_mechanical_defects(path)
    print(f"Preflight scan: {path}")
    print(f"Hits: {len(hits)}")
    for line_no, term, snippet in hits[:80]:
        print(f"  L{line_no} [{term}] {snippet!r}")
    if len(hits) > 80:
        print(f"  ... and {len(hits) - 80} more")
    print(f"Mechanical defects: {len(mechanical)}")
    for line_no, label, snippet in mechanical[:80]:
        print(f"  L{line_no} [{label}] {snippet!r}")
    if len(mechanical) > 80:
        print(f"  ... and {len(mechanical) - 80} more")
    return 1 if hits or mechanical else 0


if __name__ == "__main__":
    raise SystemExit(main())
