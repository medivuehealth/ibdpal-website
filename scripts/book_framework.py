"""Recurring chapter framework: callouts, evidence key, takeaways."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor

from book_text_cleanup import finalize_text, normalize_reference_lines

CALLOUT_COLOR = RGBColor(0x1A, 0x3A, 0x4A)

EVIDENCE_KEY = """Evidence labels used in this book:
• Established clinical practice, supported by major guidelines or strong clinical evidence.
• Promising evidence, emerging research; not established for everyone.
• Symptom-management strategy, may help symptoms without necessarily reducing intestinal inflammation.
• Individual tolerance, depends on the person rather than a universal IBD recommendation."""

SYMPTOMS_INFLAMMATION_BOX = (
    "Symptoms are not always the same as inflammation. A food that causes bloating, urgency, or "
    "discomfort is not automatically causing IBD inflammation. Likewise, feeling better after "
    "removing a food does not necessarily mean the food was driving intestinal disease activity."
)

SYMPTOMS_INFLAMMATION_MAJOR = (
    "Symptoms tell you what you notice or feel, such as pain, urgency, bloating, nausea, or stool "
    "changes. Clinical testing helps assess intestinal inflammation. The two can overlap, but they "
    "do not always move together.\n\n"
    "Symptoms alone do not prove whether intestinal inflammation has increased or decreased. "
    "Inflammation can sometimes be present even when symptoms are limited."
)

FLARE_PRIORITIES = (
    "During a flare, prioritize: adequate energy · protein · hydration · tolerated textures · "
    "clinician-directed modifications."
)

REMISSION_PRIORITIES = (
    "During remission, prioritize: dietary variety · nutrient adequacy · gradual reintroduction "
    "when appropriate · sustainable patterns · avoiding unnecessary long-term restrictions."
)

FLARE_REMISSION_COMPARE = (
    "During a flare: adequate energy, protein, hydration, tolerated textures, and "
    "clinician-directed modifications.\n\n"
    "During remission: dietary variety, nutrient adequacy, gradual reintroduction when "
    "appropriate, sustainable patterns, and avoiding unnecessary long-term restriction."
)

FLARE_REMISSION_TRANSITION = (
    "Remission is a transition toward breadth, not simply permission to eat everything. "
    "Expand variety gradually, monitor symptoms and labs, and keep your team involved when "
    "you reintroduce foods."
)

CHAPTER_TAKEAWAYS: dict[int, list[str]] = {
    1: [
        "IBD can change how you digest, absorb, and tolerate food, and symptoms do not always align neatly with laboratory markers.",
        "Flare eating and remission eating have different priorities.",
        "Symptoms after a meal do not always mean intestinal inflammation is rising.",
        "Nutrition supports care; it does not replace medical treatment.",
    ],
    6: [
        "There is no single universal IBD diet.",
        "Disease activity, location, surgery, and tolerance matter more than internet lists.",
        "Temporary texture changes during flares are tools, not lifelong bans.",
    ],
    7: [
        "During flares, enough calories, protein, and fluids come first.",
        "Gentler textures are temporary strategies, not permanent identity.",
        "Contact your team for bleeding, dehydration, obstruction signs, or inability to eat.",
    ],
    9: [
        "Food logs help when paired with labs and scopes, not guilt.",
        "Symptom triggers and inflammatory flares are related but not identical.",
    ],
    13: [
        "Deficiencies are common in IBD because of inflammation, losses, and restricted intake.",
        "Labs guide repletion; food alone may not restore low levels.",
        "Monitoring intervals should be individualized with your gastroenterology team.",
    ],
    24: [
        "FODMAP reintroduction is structured and supervised, not a pass/fail test.",
        "Pause reintroduction during active inflammation.",
        "The goal is the widest tolerated diet, not the shortest elimination list.",
    ],
    39: [
        "Sample days illustrate combinations, they are not prescriptions.",
        "Match textures to strictures, ostomy output, and active symptoms.",
        "Use foods you already tolerate as the starting point.",
    ],
    51: [
        "Your plan should update after medication changes, surgery, or sustained symptom shifts.",
        "Bring labs, logs, and specific questions to limited appointment time.",
        "Safe staples and trial foods can coexist, document both.",
    ],
}

CHAPTERS_WITH_SYMPTOMS_BOX = {6}
CHAPTERS_WITH_SYMPTOMS_CONCEPT = {9}
TRIGGERS_VS_INFLAMMATION_BOX = (
    "Food triggers and inflammatory flares are related but not identical. A symptom after a "
    "meal does not by itself prove that intestinal inflammation increased."
)
CHAPTERS_WITH_FLARE_BOX = {7, 32, 33, 34, 35, 36, 37, 38, 39}
CHAPTERS_WITH_REMISSION_BOX = {32, 33, 34, 35, 36, 37, 38}
CHAPTERS_WITH_FLARE_REMISSION_COMPARE = {8}


def _add_callout(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}\n")
    run.bold = True
    run.font.color.rgb = CALLOUT_COLOR
    run.font.size = Pt(10.5)
    run2 = p.add_run(finalize_text(body))
    run2.font.size = Pt(10.5)


def add_evidence_key_page(doc: Document, *, kindle: bool = False) -> None:
    from book_toc import add_section_heading

    add_section_heading(doc, "How to Read Evidence Labels", level=2, kindle=kindle)
    for line in EVIDENCE_KEY.split("\n"):
        doc.add_paragraph(line.strip())


def add_chapter_opening_framework(doc: Document, chapter_num: int) -> None:
    # Figure 1.2 in Chapter 1 replaces the long symptoms/inflammation callout.
    if chapter_num in CHAPTERS_WITH_SYMPTOMS_BOX:
        _add_callout(doc, "Remember", SYMPTOMS_INFLAMMATION_BOX)
    elif chapter_num in CHAPTERS_WITH_SYMPTOMS_CONCEPT:
        _add_callout(doc, "Triggers vs inflammation", TRIGGERS_VS_INFLAMMATION_BOX)
    if chapter_num in CHAPTERS_WITH_FLARE_REMISSION_COMPARE:
        _add_callout(doc, "Flare → Remission", FLARE_REMISSION_COMPARE)
        p = doc.add_paragraph()
        run = p.add_run(finalize_text(FLARE_REMISSION_TRANSITION))
        run.font.size = Pt(10.5)
    if chapter_num in CHAPTERS_WITH_FLARE_BOX:
        _add_callout(doc, "During a flare", FLARE_PRIORITIES)
    if chapter_num in CHAPTERS_WITH_REMISSION_BOX:
        _add_callout(doc, "During remission", REMISSION_PRIORITIES)


def add_chapter_closing_framework(
    doc: Document, chapter_num: int, ref_lines: list[str], *, kindle: bool = False,
) -> None:
    from book_toc import add_section_heading

    takeaways = CHAPTER_TAKEAWAYS.get(chapter_num)
    if takeaways:
        add_section_heading(doc, "Key Takeaways", kindle=kindle)
        for item in takeaways:
            doc.add_paragraph(finalize_text(item), style="List Bullet")

    if ref_lines:
        add_section_heading(doc, "References", kindle=kindle)
        for line in normalize_reference_lines(ref_lines):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(9)
