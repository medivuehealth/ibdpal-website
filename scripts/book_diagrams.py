"""Embed print-ready chapter diagrams (Figures 1.1, 1.2, 2.1)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

FIGURES_DIR_NAME = "FIGURES"
DIAGRAM_WIDTH = Inches(5.25)

CHAPTER_FIGURES: dict[int, list[tuple[str, str, str]]] = {
    1: [
        (
            "Figure_1_1_Digestive_Tract.png",
            "Figure 1.1",
            "Digestive tract overview. Crohn's disease can affect different parts of the GI tract; "
            "ulcerative colitis involves the colon and rectum.",
        ),
        (
            "Figure_1_2_Symptoms_vs_Inflammation.png",
            "Figure 1.2",
            "Symptoms versus inflammation. What you feel and what testing assesses may overlap, "
            "but they do not always move together.",
        ),
    ],
    2: [
        (
            "Figure_2_1_Disease_Location_Absorption.png",
            "Figure 2.1",
            "Disease location and absorption. Small-bowel, ileal, and colonic disease each change "
            "nutritional risk in different ways.",
        ),
    ],
}


def _figures_dir(root: Path | None) -> Path:
    base = root or Path(__file__).resolve().parent.parent
    return base / "content" / "ibd-nutrition-book" / FIGURES_DIR_NAME


def _embed_figure(
    doc: Document,
    png_path: Path,
    label: str,
    caption: str,
    credit_log: list[dict] | None,
) -> bool:
    if not png_path.is_file():
        return False
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(png_path), width=DIAGRAM_WIDTH)
        cap = doc.add_paragraph(f"{label}. {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(9)
        if credit_log is not None:
            credit_log.append(
                {
                    "source": label,
                    "file": png_path.name,
                    "alt": caption,
                    "license": "Original diagram (MediVue / Eating With IBD)",
                }
            )
        return True
    except Exception:
        return False


def add_chapter_diagrams(
    doc: Document,
    chapter_num: int,
    root=None,
    *,
    allow_text_fallback: bool = True,
    credit_log: list[dict] | None = None,
) -> int:
    """Embed PNG diagrams for chapters 1 and 2. Returns count embedded."""
    specs = CHAPTER_FIGURES.get(chapter_num, [])
    if not specs:
        return 0

    fig_dir = _figures_dir(root)
    embedded = 0
    for filename, label, caption in specs:
        if _embed_figure(doc, fig_dir / filename, label, caption, credit_log):
            embedded += 1
    return embedded
