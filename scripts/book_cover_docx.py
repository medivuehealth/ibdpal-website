"""Render hardcover front/back cover pages in Word."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from book_cover import (
    AUTHOR_BYLINE,
    AUTHOR_LINE,
    AUTHOR_NAME,
    BACK_BLURB,
    BACK_BULLETS,
    BACK_DISCLAIMER,
    BACK_HEADLINE,
    BISAC_PRIMARY,
    BISAC_SECONDARY,
    INTERIOR_PAGE_ESTIMATE,
    ISBN_13_LABEL,
    NONPROFIT_LINE,
    PRICE_CAN,
    PRICE_US,
    PUBLISHER_EMAIL,
    PUBLISHER_LEGAL,
    PUBLISHER_LOCATION,
    PUBLISHER_NAME,
    SPINE_AUTHOR,
    SPINE_TEXT,
    SPINE_WIDTH_IN_EST,
    TRIM_HEIGHT_IN,
    TRIM_WIDTH_IN,
    cover_image_path,
)
from book_manifest_data import BOOK_SUBTITLE, BOOK_TITLE
from book_standalone import ABOUT_AUTHOR

COVER_COLOR = RGBColor(0x1A, 0x3A, 0x4A)
ACCENT_COLOR = RGBColor(0xC4, 0x5C, 0x3E)
MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)


def _set_trim_size(doc: Document, width_in: float, height_in: float) -> None:
    section = doc.sections[0]
    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def _spacer(doc: Document, lines: int = 1) -> None:
    for _ in range(lines):
        doc.add_paragraph()


def _center_run(paragraph, text: str, *, size: int, bold: bool = False, italic: bool = False, color: RGBColor | None = None):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _para(doc: Document, text: str, *, size: int = 11, bold: bool = False, italic: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, color: RGBColor | None = None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def _barcode_placeholder(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ KDP barcode area, leave clear on final wraparound cover ]")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = MUTED_COLOR
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EFEFEF")
    tc_pr.append(shd)


def add_hardcover_front_cover(doc: Document, root: Path, *, standalone: bool = False) -> None:
    if standalone:
        _set_trim_size(doc, TRIM_WIDTH_IN, TRIM_HEIGHT_IN)

    label = doc.add_paragraph()
    _center_run(label, "HARDCOVER, FRONT COVER", size=9, italic=True, color=MUTED_COLOR)

    img_path = cover_image_path(root)
    if img_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Inches(4.8 if standalone else 5.5))
    else:
        _spacer(doc, 2)

    _spacer(doc, 1)
    title = doc.add_paragraph()
    _center_run(title, BOOK_TITLE.upper(), size=34 if standalone else 32, bold=True, color=COVER_COLOR)

    sub = doc.add_paragraph()
    _center_run(sub, BOOK_SUBTITLE, size=13, italic=True, color=ACCENT_COLOR)

    _spacer(doc, 2)
    author = doc.add_paragraph()
    _center_run(author, AUTHOR_LINE, size=14, bold=True, color=COVER_COLOR)

    byline = doc.add_paragraph()
    _center_run(byline, AUTHOR_BYLINE, size=10, italic=True, color=MUTED_COLOR)

    _spacer(doc, 1)
    pub = doc.add_paragraph()
    _center_run(pub, PUBLISHER_NAME, size=11, bold=True, color=COVER_COLOR)
    np = doc.add_paragraph()
    _center_run(np, NONPROFIT_LINE, size=9, italic=True, color=MUTED_COLOR)


def add_hardcover_back_cover(doc: Document, *, standalone: bool = False) -> None:
    if standalone:
        doc.add_page_break()
        _set_trim_size(doc, TRIM_WIDTH_IN, TRIM_HEIGHT_IN)

    label = doc.add_paragraph()
    _center_run(label, "HARDCOVER, BACK COVER", size=9, italic=True, color=MUTED_COLOR)

    _spacer(doc, 1)
    _para(doc, BACK_HEADLINE, size=16, bold=True, color=COVER_COLOR)

    for para in BACK_BLURB.split("\n\n"):
        _para(doc, para.strip(), size=10.5)

    _spacer(doc, 1)
    _para(doc, "Inside this book", size=12, bold=True, color=ACCENT_COLOR)
    for bullet in BACK_BULLETS:
        doc.add_paragraph(bullet, style="List Bullet")

    _spacer(doc, 1)
    _para(doc, BACK_DISCLAIMER, size=9, italic=True, color=MUTED_COLOR)

    _spacer(doc, 2)
    _barcode_placeholder(doc)

    _spacer(doc, 1)
    _para(doc, ISBN_13_LABEL, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, f"{PRICE_US}  ·  {PRICE_CAN}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, BISAC_PRIMARY, size=8, color=MUTED_COLOR)
    _para(doc, BISAC_SECONDARY, size=8, color=MUTED_COLOR)

    _spacer(doc, 1)
    _para(
        doc,
        f"{PUBLISHER_LEGAL} · {PUBLISHER_LOCATION} · {PUBLISHER_EMAIL}",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=MUTED_COLOR,
    )
    _para(doc, NONPROFIT_LINE, size=8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED_COLOR)


def add_spine_spec_page(doc: Document) -> None:
    doc.add_page_break()
    _set_trim_size(doc, TRIM_WIDTH_IN, TRIM_HEIGHT_IN)

    _para(doc, "HARDCOVER, SPINE & JACKET SPEC (FOR PRINTER)", size=9, italic=True, color=MUTED_COLOR)
    _spacer(doc, 1)
    specs = [
        f"Trim size: {TRIM_WIDTH_IN}\" × {TRIM_HEIGHT_IN}\" (confirm before press)",
        f"Estimated interior page count: ~{INTERIOR_PAGE_ESTIMATE} pages",
        f"Estimated spine width (case laminate): ~{SPINE_WIDTH_IN_EST}\", recalculate from paper stock",
        f"Spine title: {SPINE_TEXT}",
        f"Spine author/imprint: {SPINE_AUTHOR}",
        "Finish: matte laminate case wrap recommended; dust jacket optional",
        "Cover photography: Unsplash License (free to use)",
    ]
    for line in specs:
        doc.add_paragraph(line, style="List Bullet")

    _spacer(doc, 2)
    _para(doc, "Flat spine text (center, bottom to top on jacket):", size=10, bold=True)
    _para(
        doc,
        f"{SPINE_TEXT}    ·    {SPINE_AUTHOR}",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=COVER_COLOR,
    )

    _spacer(doc, 2)
    _para(
        doc,
        "Dust-jacket spread order (outside flat): Back cover | Spine | Front cover, "
        "add bleed per printer template (typically 0.125\").",
        size=9,
        italic=True,
        color=MUTED_COLOR,
    )


def add_about_author_page(
    doc: Document, root: Path, *, kindle: bool = False, nav=None,
) -> None:
    from book_toc import add_back_matter_heading

    doc.add_page_break()
    add_back_matter_heading(doc, "About the Author", kindle=kindle, nav=nav)

    name_img = root / "assets" / "founder-name.png"
    if name_img.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(name_img), width=Inches(4.5))

    author_head = doc.add_paragraph()
    _center_run(author_head, AUTHOR_NAME, size=16, bold=True, color=COVER_COLOR)
    role = doc.add_paragraph()
    _center_run(role, AUTHOR_BYLINE, size=11, italic=True, color=MUTED_COLOR)

    _spacer(doc, 1)
    for para in ABOUT_AUTHOR.split("\n\n"):
        _para(doc, para.strip(), size=11)


def add_interior_half_title(doc: Document) -> None:
    doc.add_page_break()
    _spacer(doc, 6)
    title = doc.add_paragraph()
    _center_run(title, BOOK_TITLE, size=28, bold=True)
    sub = doc.add_paragraph()
    _center_run(sub, BOOK_SUBTITLE, size=13, italic=True, color=MUTED_COLOR)
    _spacer(doc, 2)
    author = doc.add_paragraph()
    _center_run(author, AUTHOR_LINE, size=12, bold=True)
    _spacer(doc, 4)
    pub = doc.add_paragraph()
    _center_run(pub, PUBLISHER_NAME, size=11, color=MUTED_COLOR)


def build_standalone_cover_doc(root: Path, output: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    add_hardcover_front_cover(doc, root, standalone=True)
    add_hardcover_back_cover(doc, standalone=True)
    add_spine_spec_page(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
