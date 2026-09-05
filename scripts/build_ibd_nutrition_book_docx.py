#!/usr/bin/env python3
"""Compile nutrition corpus into a standalone book Word document."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(Path(__file__).resolve().parent))

from book_manifest_data import (  # noqa: E402
    BOOK_SUBTITLE,
    BOOK_TITLE,
    CHAPTERS,
    NEW_CHAPTER_STUBS,
    PARTS,
)
from book_standalone import (  # noqa: E402
    ABOUT_BOOK,
    AUTOIMMUNE_NOTE,
    CONCLUSION,
    DRI_SOURCE_NOTE,
    EDUCATIONAL_DISCLAIMER,
    FOOD_INTRO,
    HOW_TO_USE,
    IMAGE_LICENSE_NOTE,
    STANDALONE_PREFACE,
    START_HERE,
    heading_should_skip,
    sanitize_book_text,
    sanitize_title,
)
from book_images import extract_licensed_images  # noqa: E402
from book_cover_docx import (  # noqa: E402
    add_about_author_page,
    add_hardcover_back_cover,
    add_hardcover_front_cover,
    add_interior_half_title,
    build_standalone_cover_doc,
)
from book_cover import AUTHOR_NAME  # noqa: E402
from book_citations import CHAPTER_REFERENCES, CitationRegistry  # noqa: E402
from book_framework import (  # noqa: E402
    add_chapter_closing_framework,
    add_chapter_opening_framework,
    add_evidence_key_page,
)
from book_chapter_filters import (  # noqa: E402
    GlobalDeficiencyStackTracker,
    filter_blocks_for_chapter,
)
from book_merge import (  # noqa: E402
    ChapterDeduper,
    GlobalBoilerplateTracker,
    GlobalDeduper,
    demote_headings_for_merge,
    filter_blocks_for_merge,
    is_food_reference_chapter,
    repair_blocks,
)
from book_text_cleanup import clean_blocks, finalize_text, format_book_heading  # noqa: E402
from book_food_reference import blocks_to_food_entry  # noqa: E402
from book_prose_cleanup import GlobalRepeatTracker, filter_chapter_sections, heading_starts_skip_section, strip_boilerplate_tail, norm_heading
from book_chapter_rewrites import get_chapter_rewrite, get_chapter_prepend  # noqa: E402
from book_diagrams import add_chapter_diagrams  # noqa: E402
from book_appendices import (  # noqa: E402
    add_abbreviations,
    add_food_index,
    add_glossary,
    add_index,
    add_red_flags_page,
)

OUTPUT = ROOT / "content" / "ibd-nutrition-book" / "Eating_With_IBD_Interior.docx"
DRAFT_OUTPUT = ROOT / "content" / "ibd-nutrition-book" / "Eating_With_IBD_DRAFT.docx"
KINDLE_OUTPUT = ROOT / "content" / "ibd-nutrition-book" / "KINDLE" / "Eating_With_IBD_Kindle.docx"
COVER_OUTPUT = ROOT / "content" / "ibd-nutrition-book" / "Eating_With_IBD_HARDCOVER_COVERS.docx"
DRI_JSON = ROOT / "data" / "nutrition-dri-baselines.json"
FOOD_JSON = ROOT / "data" / "nutrition-food-sources.json"
from book_toc import (  # noqa: E402
    KindleNavRegistry,
    add_back_matter_heading,
    add_book_heading,
    add_chapter_table_of_contents,
    add_front_matter_heading,
    add_section_heading,
    update_docx_fields,
    verify_print_toc,
)

MAX_IMAGES_FOOD_CHAPTER = 12
MAX_IMAGES_DEFAULT_CHAPTER = 0
MAX_IMAGES_PER_SOURCE_FOOD = 1
IMAGE_WIDTH = Inches(5.25)


@dataclass
class BuildConfig:
    edition: str  # interior | draft | kindle
    output: Path
    include_hardcover_in_interior: bool = False
    include_compile_report: bool = False
    allow_diagram_text_fallback: bool = False
    write_cover_reference: bool = True


def edition_config(edition: str) -> BuildConfig:
    if edition == "draft":
        return BuildConfig(
            edition="draft",
            output=DRAFT_OUTPUT,
            include_hardcover_in_interior=True,
            include_compile_report=True,
            allow_diagram_text_fallback=True,
            write_cover_reference=True,
        )
    if edition == "kindle":
        return BuildConfig(
            edition="kindle",
            output=KINDLE_OUTPUT,
            include_hardcover_in_interior=False,
            include_compile_report=False,
            allow_diagram_text_fallback=False,
            write_cover_reference=False,
        )
    return BuildConfig(
        edition="interior",
        output=OUTPUT,
        include_hardcover_in_interior=False,
        include_compile_report=False,
        allow_diagram_text_fallback=False,
        write_cover_reference=True,
    )

SKIP_CLASSES = {
    "blog-medical-review",
    "blog-edu-disclaimer",
    "community-edu-disclaimer",
    "blog-figure-grid",
    "blog-photo-credit",
    "blog-vote",
    "blog-footer",
    "blog-share",
    "seo-guide-keywords",
    "blog-back",
    "seo-related-reading",
}

SKIP_HEADINGS = set()  # legacy; use heading_should_skip from book_standalone


def site_path_to_file(site_path: str) -> Path:
    p = site_path.strip("/")
    if p.endswith(".html"):
        p = p[:-5]
    if p.startswith("blog/"):
        return ROOT / "blogs" / (p.split("/", 1)[1] + ".html")
    if p.startswith("guides/"):
        return ROOT / "guides" / (p.split("/", 1)[1] + ".html")
    return ROOT / (p + ".html")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def should_skip_element(el: Tag) -> bool:
    classes = el.get("class") or []
    if any(c in SKIP_CLASSES for c in classes):
        return True
    if "seo-landing__list" in classes:
        return True
    for parent in el.parents:
        if not isinstance(parent, Tag):
            break
        pclasses = parent.get("class") or []
        if "seo-related-reading" in pclasses:
            return True
    if el.name in ("script", "style", "figure", "img"):
        return True
    return False


def heading_is_skip(text: str) -> bool:
    return heading_should_skip(text)


def extract_blocks_from_soup(
    soup: BeautifulSoup,
    *,
    max_images: int = 3,
    food_only_images: bool = True,
) -> list[tuple[str, str | list[str]]]:
    """Return list of (kind, content) where kind is heading|paragraph|list."""
    article = soup.select_one("div.blog-content")
    if article is None:
        article = soup.select_one("article.support-section.seo-landing")
    if article is None:
        article = soup.select_one("article")
    if article is None:
        return []

    blocks: list[tuple[str, str | list[str]]] = []
    skip_until_next_heading = False
    food_mode = food_only_images

    for el in article.descendants:
        if not isinstance(el, Tag):
            continue
        if should_skip_element(el):
            continue
        if el.name in ("h1", "h2", "h3", "h4"):
            text = normalize_text(el.get_text())
            if not text:
                continue
            clean = sanitize_book_text(text)
            if clean is None:
                continue
            if heading_is_skip(clean):
                skip_until_next_heading = True
                continue
            if heading_starts_skip_section(clean, food_mode=food_mode):
                skip_until_next_heading = True
                continue
            skip_until_next_heading = False
            level = el.name
            blocks.append((f"heading_{level}", clean))
        elif el.name == "p" and not skip_until_next_heading:
            if el.find_parent(["h1", "h2", "h3", "h4", "li", "figure"]):
                continue
            classes = el.get("class") or []
            if any(c in SKIP_CLASSES for c in classes):
                continue
            text = normalize_text(el.get_text())
            if not text or text.lower().startswith("related reading:"):
                continue
            clean = sanitize_book_text(text)
            if clean is None:
                continue
            blocks.append(("paragraph", clean))
        elif el.name in ("ul", "ol") and not skip_until_next_heading:
            if el.find_parent(["ul", "ol"]):
                continue
            items = []
            for li in el.find_all("li", recursive=False):
                t = normalize_text(li.get_text())
                if not t:
                    continue
                clean = sanitize_book_text(t)
                if clean:
                    items.append(clean)
            if items:
                blocks.append(("list", items))

    for image in extract_licensed_images(
        soup, ROOT, max_images=max_images, food_only=food_only_images,
    ):
        blocks.append(("image", image))

    return dedupe_blocks(blocks)


def dedupe_blocks(blocks: list[tuple[str, str | list[str]]]) -> list[tuple[str, str | list[str]]]:
    seen_paragraphs: set[str] = set()
    out: list[tuple[str, str | list[str]]] = []
    for kind, content in blocks:
        if kind == "paragraph" and isinstance(content, str):
            key = content[:120]
            if key in seen_paragraphs:
                continue
            seen_paragraphs.add(key)
        out.append((kind, content))
    return out


def extract_source(
    site_path: str,
    *,
    max_images: int = 3,
    food_only_images: bool = True,
) -> tuple[str, list[tuple[str, str | list[str]]]]:
    fpath = site_path_to_file(site_path)
    if not fpath.exists():
        return site_path, [("paragraph", f"[Source missing: {site_path}]")]
    html = fpath.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    title = soup.title.string if soup.title and soup.title.string else site_path
    title = normalize_text(re.sub(r"\s*\|\s*IBDPal.*", "", title))
    title = sanitize_title(title)
    return title, extract_blocks_from_soup(
        soup, max_images=max_images, food_only_images=food_only_images,
    )


def add_heading(doc: Document, text: str, level: int = 1, *, kindle: bool = False) -> None:
    """Front-matter subheads excluded from print auto-TOC and Kindle NCX."""
    normalized = format_book_heading(text)
    if level == 1:
        add_front_matter_heading(doc, normalized, kindle=kindle)
    else:
        add_section_heading(doc, normalized, level=level, kindle=kindle)


def add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(finalize_text(text))
    if italic:
        run.italic = True


def add_image_block(
    doc: Document,
    image: dict,
    embedded_paths: set[str],
    credit_log: list[dict],
    source_title: str,
) -> bool:
    path = Path(image["path"])
    if not path.is_file():
        return False
    key = str(path.resolve())
    if key in embedded_paths:
        return False
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=IMAGE_WIDTH)
        alt = str(image.get("alt") or "").strip()
        if not alt or alt.lower() in {"image", "photo", "food"}:
            alt = f"Illustration from {source_title}" if source_title else "Educational food photograph"
        credit = str(image.get("credit") or "").strip()
        if alt:
            cap = doc.add_paragraph(alt)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)
        embedded_paths.add(key)
        credit_log.append(
            {
                "source": source_title,
                "file": path.name,
                "alt": alt,
                "license": credit or "Free use",
            }
        )
        return True
    except Exception:
        return False


def add_blocks(
    doc: Document,
    blocks: list[tuple[str, str | list[str] | dict]],
    skip_h1: bool = True,
    source_title: str = "",
    embedded_paths: set[str] | None = None,
    credit_log: list[dict] | None = None,
    max_images: int | None = None,
    *,
    kindle: bool = False,
) -> int:
    images_added = 0
    paths = embedded_paths if embedded_paths is not None else set()
    credits = credit_log if credit_log is not None else []
    for kind, content in blocks:
        if kind.startswith("heading_"):
            level = int(kind[-1])
            if skip_h1 and level == 1:
                continue
            add_section_heading(doc, str(content), level=min(level, 3), kindle=kindle)
        elif kind == "paragraph":
            add_paragraph(doc, str(content))
        elif kind == "list" and isinstance(content, list):
            for item in content:
                doc.add_paragraph(finalize_text(str(item)), style="List Bullet")
        elif kind == "image" and isinstance(content, dict):
            if max_images is not None and images_added >= max_images:
                continue
            if add_image_block(doc, content, paths, credits, source_title):
                images_added += 1
    return images_added


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def append_dri_tables(doc: Document, dri: dict, *, kindle: bool = False) -> None:
    add_section_heading(doc, "Dietary Reference Intake Baselines", level=2, kindle=kindle)
    add_paragraph(
        doc,
        f"{DRI_SOURCE_NOTE} Values below are population reference ranges from the National "
        "Academies Dietary Reference Intakes; estimated energy ranges follow National Academies "
        "Estimated Energy Requirement (EER) data and are not RDA or AI values. Individual targets "
        "may differ with IBD activity, malabsorption, or treatment.",
    )
    for profile in dri.get("profiles", []):
        add_section_heading(
            doc,
            profile.get("label", profile.get("id", "Profile")),
            level=3,
            kindle=kindle,
        )
        for section_key, section_title in (("macros", "Macronutrients"), ("micros", "Micronutrients")):
            rows = profile.get(section_key, [])
            if not rows:
                continue
            add_section_heading(doc, section_title, level=4, kindle=kindle)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Nutrient"
            hdr[1].text = "Baseline"
            hdr[2].text = "IBD note"
            for row in rows:
                cells = table.add_row().cells
                cells[0].text = row.get("nutrient", "")
                cells[1].text = row.get("baseline", "")
                note = row.get("note", "")
                if note:
                    cells[2].text = note
            doc.add_paragraph()


def append_food_source(doc: Document, food_data: dict, nutrient_ids: list[str]) -> None:
    nutrients = {n["id"]: n for n in food_data.get("nutrients", [])}
    for nid in nutrient_ids:
        n = nutrients.get(nid)
        if not n:
            add_paragraph(doc, f"[Food source data missing: {nid}]")
            continue
        add_heading(doc, f"Food sources: {n.get('name', nid)}", level=3)
        if n.get("ibdNote"):
            note = sanitize_book_text(n["ibdNote"])
            if note:
                add_paragraph(doc, note)
        foods = n.get("foods", [])
        if foods:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Food"
            hdr[1].text = "Serving"
            hdr[2].text = "Amount"
            for f in foods[:12]:
                cells = table.add_row().cells
                cells[0].text = f.get("food", "")
                cells[1].text = f.get("serving", "")
                cells[2].text = f.get("amount", "")
            doc.add_paragraph()


def add_front_matter(doc: Document, *, kindle: bool = False, nav: KindleNavRegistry | None = None) -> None:
    add_heading(doc, "Copyright and Educational Disclaimer", level=1, kindle=kindle)
    add_paragraph(
        doc,
        f"© 2026 {AUTHOR_NAME} and MediVue. All rights reserved.",
    )
    for para in EDUCATIONAL_DISCLAIMER.split("\n\n"):
        add_paragraph(doc, para.strip())
    add_paragraph(doc, IMAGE_LICENSE_NOTE)

    add_page_break(doc)

    add_front_matter_heading(doc, "Preface", kindle=kindle)
    for para in STANDALONE_PREFACE.split("\n\n"):
        add_paragraph(doc, para.strip())

    add_heading(doc, "How to Use This Book", level=2, kindle=kindle)
    for para in HOW_TO_USE.split("\n\n"):
        add_paragraph(doc, para.strip())

    add_heading(doc, "A Note on Autoimmune Overlap", level=2, kindle=kindle)
    add_paragraph(doc, AUTOIMMUNE_NOTE)

    add_evidence_key_page(doc, kindle=kindle)

    add_page_break(doc)
    add_front_matter_heading(doc, "Table of Contents", kindle=kindle)
    add_chapter_table_of_contents(
        doc, kindle=kindle, chapters=CHAPTERS, parts=PARTS, nav=nav,
    )
    add_page_break(doc)

    add_front_matter_heading(doc, "Start Here", kindle=kindle)
    add_paragraph(
        doc,
        "This book is long by design. Use the paths below to jump to what matters most right now. "
        "Chapter numbers match the table of contents.",
    )
    for line in START_HERE.split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line, style="List Bullet")
    add_page_break(doc)


def add_illustration_credits_appendix(
    doc: Document,
    credit_log: list[dict],
    *,
    kindle: bool = False,
    nav: KindleNavRegistry | None = None,
) -> None:
    add_page_break(doc)
    add_back_matter_heading(doc, "Illustration Credits", kindle=kindle, nav=nav)
    add_paragraph(doc, IMAGE_LICENSE_NOTE)
    add_paragraph(
        doc,
        "The table lists photographs embedded in this edition. Typeset callout boxes appear "
        "throughout the manuscript. Third-party photo licenses are documented in "
        "blogs/assets/IMAGE_CREDITS.md.",
    )
    if not credit_log:
        add_paragraph(doc, "No illustrations were embedded in this build.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "File"
    hdr[2].text = "Description"
    hdr[3].text = "License"
    for row in credit_log:
        cells = table.add_row().cells
        cells[0].text = row.get("source", "")
        cells[1].text = row.get("file", "")
        cells[2].text = row.get("alt", "")
        cells[3].text = row.get("license", "")


def add_review_section(doc: Document, stats: dict) -> None:
    add_page_break(doc)
    add_heading(doc, "Manuscript Compile Report", level=1)
    add_paragraph(
        doc,
        f"Generated: {date.today().isoformat()}. Editorial revision pipeline applied.",
    )
    add_heading(doc, "Statistics", level=2)
    skip_stats = {"embedded_image_paths", "image_credits"}
    for key, val in stats.items():
        if key in skip_stats:
            continue
        doc.add_paragraph(f"{key}: {val}", style="List Bullet")

    add_heading(doc, "Remaining before publication", level=2)
    review_items = [
        "Optional final line edit for one unified author voice",
        "Replace diagram placeholders with designed figures before print",
        "Verify chapter references against latest guideline editions",
        "Refresh index page numbers after final layout",
        "Assign ISBN and finalize copyright / imprint page",
        "Generate and inspect final print PDF page by page",
    ]
    for item in review_items:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Missing or failed sources", level=2)
    missing = stats.get("missing_sources", [])
    if missing:
        for m in missing:
            doc.add_paragraph(m, style="List Bullet")
    else:
        add_paragraph(doc, "All manifest source paths resolved successfully.")


def compile_chapter_blocks(
    chapter: dict,
    deduper: GlobalDeduper,
    global_boilerplate: GlobalBoilerplateTracker,
    global_repeat: GlobalRepeatTracker,
    deficiency_stack: GlobalDeficiencyStackTracker,
) -> tuple[list[tuple[str, str | list[str] | dict]], int, list[str]]:
    """Return merged blocks, source count, and missing paths for a chapter."""
    ch_num = chapter["num"]
    missing: list[str] = []
    rewrite = get_chapter_rewrite(ch_num)
    if rewrite:
        return clean_blocks(list(rewrite)), 0, missing

    if chapter.get("stub") and not chapter.get("rewrite"):
        stub = NEW_CHAPTER_STUBS.get(ch_num, "[Draft stub, expand in rewrite pass.]")
        blocks: list[tuple[str, str | list[str] | dict]] = []
        for para in stub.split("\n\n"):
            text = para.strip()
            if text.startswith("**") and text.endswith("**"):
                blocks.append(("heading_h3", text.strip("*")))
            elif text.startswith("**"):
                parts = text.split("**")
                if len(parts) >= 3:
                    blocks.append(("heading_h3", parts[1]))
                    blocks.append(("paragraph", parts[2].strip(": ").strip()))
                else:
                    blocks.append(("paragraph", text))
            else:
                blocks.append(("paragraph", text))
        return blocks, 0, missing

    ch_num = chapter["num"]
    food_mode = is_food_reference_chapter(chapter)
    merged: list[tuple[str, str | list[str] | dict]] = []
    source_count = 0
    first_source = True
    chapter_deduper = ChapterDeduper()

    for src in chapter.get("sources", []):
        if deduper.should_skip_source(src, allow_reuse=food_mode):
            continue
        title, raw_blocks = extract_source(
            src,
            max_images=MAX_IMAGES_PER_SOURCE_FOOD if food_mode else 0,
            food_only_images=food_mode,
        )
        if raw_blocks == [("paragraph", f"[Source missing: {src}]")]:
            missing.append(src)
            continue
        deduper.mark_source_used(src)
        source_count += 1

        if food_mode:
            entry_blocks = blocks_to_food_entry(title, raw_blocks)  # type: ignore[arg-type]
            filtered = filter_blocks_for_merge(
                entry_blocks, chapter_deduper, food_mode=True,
                global_boilerplate=global_boilerplate, global_repeat=global_repeat,
            )
        else:
            filtered = filter_blocks_for_merge(
                raw_blocks, chapter_deduper, food_mode=False,
                global_boilerplate=global_boilerplate, global_repeat=global_repeat,
            )
            if not first_source:
                filtered = demote_headings_for_merge(filtered)

        filtered = filter_blocks_for_chapter(ch_num, filtered)
        merged.extend(filtered)
        first_source = False

    merged = repair_blocks(merged)
    prepend = get_chapter_prepend(ch_num)
    if prepend:
        merged = list(prepend) + merged
    has_intro = bool(chapter.get("intro")) or bool(chapter.get("rewrite"))
    merged = filter_blocks_for_chapter(
        ch_num,
        merged,
        deficiency_stack=deficiency_stack,
        drop_lede=True,
        has_intro=has_intro,
    )
    # Food entries repeat section labels (During a flare, Nutrition, …) by design.
    if not food_mode:
        merged = filter_chapter_sections(merged)
    merged = strip_boilerplate_tail(merged)
    merged = clean_blocks(merged)
    return merged, source_count, missing


def build_document(config: BuildConfig) -> dict:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if config.include_hardcover_in_interior:
        add_hardcover_front_cover(doc, ROOT, standalone=False)
        add_interior_half_title(doc)
    elif config.edition == "kindle":
        add_interior_half_title(doc)
    else:
        add_interior_half_title(doc)

    kindle = config.edition == "kindle"
    nav = KindleNavRegistry() if kindle else None

    add_front_matter(doc, kindle=kindle, nav=nav)

    dri = load_json(DRI_JSON)
    food_data = load_json(FOOD_JSON)
    deduper = GlobalDeduper()
    global_boilerplate = GlobalBoilerplateTracker()
    deficiency_stack = GlobalDeficiencyStackTracker()
    global_repeat = GlobalRepeatTracker()
    citations = CitationRegistry()

    stats = {
        "chapters": len(CHAPTERS),
        "parts": len(PARTS),
        "words": 0,
        "sources_compiled": 0,
        "images_embedded": 0,
        "images_skipped_duplicate": 0,
        "missing_sources": [],
        "embedded_image_paths": set(),
        "image_credits": [],
    }

    current_part = None
    for chapter in CHAPTERS:
        part_num = chapter["part"]
        if part_num != current_part:
            current_part = part_num
            part_meta = next(p for p in PARTS if p["num"] == part_num)
            add_page_break(doc)
            add_book_heading(
                doc,
                f"Part {part_num}: {part_meta['title']}",
                part=True,
                kindle=kindle,
                nav=nav,
                part_num=part_num,
            )

        ch_num = chapter["num"]
        add_book_heading(
            doc,
            f"Chapter {ch_num}: {chapter['title']}",
            part=False,
            kindle=kindle,
            nav=nav,
            chapter_num=ch_num,
        )

        if chapter.get("intro") and not chapter.get("rewrite"):
            add_paragraph(doc, chapter["intro"], italic=True)

        add_chapter_opening_framework(doc, ch_num)
        add_chapter_diagrams(
            doc,
            ch_num,
            ROOT,
            allow_text_fallback=config.allow_diagram_text_fallback,
            credit_log=stats["image_credits"],
        )

        blocks, src_count, missing = compile_chapter_blocks(
            chapter, deduper, global_boilerplate, global_repeat, deficiency_stack,
        )
        stats["sources_compiled"] += src_count
        stats["missing_sources"].extend(missing)

        img_cap = MAX_IMAGES_FOOD_CHAPTER if is_food_reference_chapter(chapter) else MAX_IMAGES_DEFAULT_CHAPTER
        # Food chapters may legitimately reuse a photo (e.g. potato in starches and vegetables).
        chapter_image_paths: set[str] = set() if is_food_reference_chapter(chapter) else stats["embedded_image_paths"]
        add_blocks(
            doc,
            blocks,
            skip_h1=True,
            source_title=chapter["title"],
            embedded_paths=chapter_image_paths,
            credit_log=stats["image_credits"],
            max_images=img_cap,
            kindle=kindle,
        )
        if is_food_reference_chapter(chapter):
            stats["embedded_image_paths"].update(chapter_image_paths)

        if chapter.get("append_dri"):
            append_dri_tables(doc, dri, kindle=kindle)

        if chapter.get("append_food_sources"):
            append_food_source(doc, food_data, chapter["append_food_sources"])

        ref_ids = citations.assign_chapter(ch_num, CHAPTER_REFERENCES.get(ch_num, [1, 2]))
        ref_lines = citations.chapter_reference_lines(ch_num)
        add_chapter_closing_framework(doc, ch_num, ref_lines, kindle=kindle)

        if chapter.get("bridge"):
            add_paragraph(doc, chapter["bridge"], italic=True)

    # Back matter
    add_page_break(doc)
    add_back_matter_heading(
        doc, "Conclusion: Eating With IBD Over Time", kindle=kindle, nav=nav,
    )
    for para in CONCLUSION.split("\n\n"):
        add_paragraph(doc, para.strip())

    add_about_author_page(doc, ROOT, kindle=kindle, nav=nav)

    add_heading(doc, "About This Book", level=2, kindle=kindle)
    for para in ABOUT_BOOK.split("\n\n"):
        add_paragraph(doc, para.strip())

    add_heading(doc, "Appendix: Food Data Notes", level=2, kindle=kindle)
    add_paragraph(doc, FOOD_INTRO)

    add_red_flags_page(doc, kindle=kindle, nav=nav)
    add_glossary(doc, kindle=kindle, nav=nav)
    add_abbreviations(doc, kindle=kindle, nav=nav)
    add_food_index(doc, kindle=kindle, nav=nav)
    if config.edition != "kindle":
        add_index(doc)
    else:
        add_index(doc, kindle=True, nav=nav)

    stats["images_embedded"] = len(stats["embedded_image_paths"])
    add_illustration_credits_appendix(
        doc, stats["image_credits"], kindle=kindle, nav=nav,
    )

    if config.include_hardcover_in_interior:
        add_page_break(doc)
        add_hardcover_back_cover(doc, standalone=False)

    # Word count from document text
    full_text = "\n".join(p.text for p in doc.paragraphs)
    stats["words"] = len(full_text.split())
    stats["estimated_pages"] = f"~{max(1, stats['words'] // 300)} pages at ~300 words/page"
    stats["edition"] = config.edition

    if config.include_compile_report:
        add_review_section(doc, stats)

    config.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(config.output))
    fields_updated = False if kindle else update_docx_fields(str(config.output))
    stats["toc_fields_updated"] = fields_updated
    toc_errors = verify_print_toc(config.output, kindle=kindle)
    if toc_errors:
        stats["toc_verify_errors"] = toc_errors
        for err in toc_errors:
            print(f"TOC verify: {err}")
    stats["output"] = str(config.output)
    return stats


def main() -> int:
    parser = argparse.ArgumentParser(description="Build Eating With IBD Word editions")
    parser.add_argument(
        "--edition",
        choices=("interior", "draft", "kindle", "all"),
        default="interior",
        help="interior=KDP print; draft=dev; kindle=reflowable; all=interior+kindle",
    )
    parser.add_argument(
        "--figures",
        action="store_true",
        help="Generate diagram PNG/SVG before building",
    )
    parser.add_argument(
        "--no-archive",
        action="store_true",
        help="Skip copying built files to MASTER_CONTENT/ARCHIVE",
    )
    args = parser.parse_args()

    from book_publication import archive_locked_masters, run_preflight, validate_chapter_references

    ref_errors = validate_chapter_references()
    if ref_errors:
        print("Reference validation failed:")
        for err in ref_errors:
            print(f"  - {err}")
        return 1

    from audit_nutrition_values import main as audit_nutrition_main

    if audit_nutrition_main() != 0:
        return 1

    if args.figures:
        from generate_book_figures import main as gen_figures

        if gen_figures() != 0:
            return 1

    editions = ("interior", "kindle") if args.edition == "all" else (args.edition,)
    last_stats: dict | None = None
    built_paths: list[Path] = []

    for edition in editions:
        config = edition_config(edition)
        stats = build_document(config)
        last_stats = stats
        if config.write_cover_reference and edition == "interior":
            build_standalone_cover_doc(ROOT, COVER_OUTPUT)
            print(f"Wrote {COVER_OUTPUT}")
        print(f"Wrote {config.output}")
        print(f"Edition: {edition}")
        print(f"Words: {stats['words']:,}")
        print(f"Estimated length: {stats['estimated_pages']}")
        print(f"Sources compiled: {stats['sources_compiled']}")
        print(f"Images embedded: {stats['images_embedded']}")
        if stats.get("toc_fields_updated"):
            print("TOC page numbers updated via Word.")
        elif stats.get("toc_fields_updated") is False:
            print("TOC page numbers: open in Word and press Ctrl+A, F9 to refresh (or install pywin32).")
        if stats["missing_sources"]:
            print(f"Missing sources ({len(stats['missing_sources'])}):")
            for m in stats["missing_sources"]:
                print(f"  - {m}")
            return 1
        built_paths.append(config.output)

    if not args.no_archive and "interior" in editions:
        archived = archive_locked_masters()
        for path in archived:
            print(f"Archived: {path}")

    preflight_paths = [p for p in built_paths if p.name.endswith(".docx")]
    if preflight_paths:
        print("\n--- Preflight ---")
        if run_preflight(preflight_paths):
            return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
